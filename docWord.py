from matplotlib import pyplot as plt
import pandas as pd
from docx import Document
from docx.shared import Inches, Cm   
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from openpyxl.drawing.image import Image as xlImage
from openpyxl.styles import Font, Alignment, Border, Side
import six
import numpy as np
from PyQt5.QtWidgets import QScrollArea, QLabel, QWidget, QMainWindow, QVBoxLayout
from PyQt5.QtGui import QIcon, QPixmap
from PyQt5.QtCore import Qt, QSize
from PyQt5.QtWidgets import QMainWindow, QWidget, QVBoxLayout, QListWidget, QListWidgetItem, QPushButton, QLabel, QScrollArea
from PyQt5.QtGui import QPixmap
from PyQt5.QtCore import Qt
import tempfile
import os

from funcCapCargAoki import resultAoki, paFinalAoki
from funcDecQuar import resulDecQuar, cotasPonta, paFinalDecQuar
from resultNotavel import resulNotaveis, minAokiDec



from PIL import Image





'''
# Chamando a função com os dados apropriados
tipoEstaca = 'HCM'      # HCM, Escavada, Raiz
diametroEst = 0.3       # m
alturaBloco = 0.45      # m
cargaAdmEsperada = 170  # kN
listaNspt = [0, 1, 1, 3, 6, 10, 11, 16, 30, 50, 50, 50, 50, 50, 50]
listaTipoSolo = [23, 23, 23, 23, 32, 32, 32, 32, 32, 12, 12, 12, 12, 12, 12]
fileName = 'arqWord'
'''

def render_mpl_table(data, col_width=3.0, row_height=0.625, font_size=14, header_font_size=7,
                     header_color='#40466e', row_colors=['#f1f1f2', 'w'], edge_color='k',
                     bbox=[0, 0, 1, 1], header_columns=0, ax=None, **kwargs):
    if ax is None:
        size = (np.array(data.shape[::-1]) + np.array([0, 1])) * np.array([col_width, row_height])
        fig, ax = plt.subplots(figsize=size)
        ax.axis('off')

    colLabels = [f"\n{col.replace(' ', '\n')}\n" for col in data.columns]
    mpl_table = ax.table(cellText=data.values, bbox=bbox, colLabels=colLabels, **kwargs)

    mpl_table.auto_set_font_size(False)
    mpl_table.set_fontsize(font_size)

    for k, cell in six.iteritems(mpl_table._cells):
        cell.set_edgecolor(edge_color)
        if k[0] == 0:
            cell.set_text_props(weight='bold', color='w', size=header_font_size)
            cell.set_facecolor(header_color)
        else:
            cell.set_facecolor(row_colors[k[0] % len(row_colors)])
            cell.set_text_props(size=font_size)
    return ax

def geracaoGraficos(listaTipoSolo, tipoEstaca, listaNspt, diametroEst, alturaBloco, cargaAdmEsperada):
    # Suponho que as funções minAokiDec, cotasPonta, paFinalAoki, paFinalDecQuar, resultAoki,
    # resulDecQuar e resulNotaveis estejam definidas em algum lugar do seu código
    menorAokiDec = minAokiDec(listaTipoSolo, tipoEstaca, listaNspt, diametroEst)
    cota = cotasPonta(listaTipoSolo)
    aoki = paFinalAoki(listaTipoSolo, tipoEstaca, listaNspt, diametroEst)
    decQuar = paFinalDecQuar(listaTipoSolo, tipoEstaca, listaNspt, diametroEst)

    topoBloco = -1 + alturaBloco
    fundoBloco = -1

    df_aoki = pd.DataFrame(resultAoki(listaTipoSolo, tipoEstaca, listaNspt, diametroEst))
    df_decQuar = pd.DataFrame(resulDecQuar(listaTipoSolo, tipoEstaca, listaNspt, diametroEst))
    df_notaveis = pd.DataFrame(resulNotaveis(listaTipoSolo, tipoEstaca, listaNspt, diametroEst))

    
    render_mpl_table(df_aoki, header_columns=0, col_width=3.0)
    plt.savefig('tabela_aoki.png', bbox_inches='tight', dpi=100)  # Reduzindo o DPI para 100
    render_mpl_table(df_decQuar, header_columns=0, col_width=3.0)
    plt.savefig('tabela_decQuar.png', bbox_inches='tight', dpi=100)  # Reduzindo o DPI para 100
    render_mpl_table(df_notaveis, header_columns=0, col_width=3.0)
    plt.savefig('tabela_notaveis.png', bbox_inches='tight', dpi=100)  # Reduzindo o DPI para 100

    plt.figure(figsize=(6, 8))  # Reduzindo o tamanho da figura
    plt.plot(menorAokiDec, cota, color='green', linestyle='-', linewidth=2, label='Menor entre Aoki e DecQua')
    plt.scatter(menorAokiDec, cota, color='red', marker='.', linewidth=3, s=100)
    plt.axvline(cargaAdmEsperada, color='purple', linestyle='--', linewidth=3, label='Carga Adm Esperada')
    plt.axhline(0, color='k', linestyle='dotted', linewidth=2, label='N.T')
    plt.axhline(topoBloco, color='k', linestyle='dotted', linewidth=1, label='Topo do Bloco')
    plt.axhline(fundoBloco, color='k', linestyle='dotted', linewidth=1, label='Fundo do Bloco')
    plt.xlabel('Carga Admissível (kN)', fontsize=12)
    plt.ylabel('Cotas (m)', fontsize=12)
    plt.title('Carga Adm Final (kN)', fontsize=14)
    plt.grid(True)
    plt.legend()
    plt.tight_layout()
    plt.savefig('grafico1.png', bbox_inches='tight', dpi=100)  # Reduzindo o DPI para 100

    plt.figure(figsize=(6, 8))  # Reduzindo o tamanho da figura
    plt.plot(decQuar, cota, color='green', linestyle='-', linewidth=2, label='Décourt-Quaresma')
    plt.scatter(decQuar, cota, color='red', marker='.', linewidth=3, s=100)
    plt.axvline(cargaAdmEsperada, color='purple', linestyle='--', linewidth=3, label='Carga Adm Esperada')
    plt.axhline(0, color='k', linestyle='dotted', linewidth=2, label='N.T')
    plt.axhline(topoBloco, color='k', linestyle='dotted', linewidth=1, label='Topo do Bloco')
    plt.axhline(fundoBloco, color='k', linestyle='dotted', linewidth=1, label='Fundo do Bloco')
    plt.xlabel('Carga Admissível (kN)', fontsize=12)
    plt.ylabel('Cotas (m)', fontsize=12)
    plt.title('Carga Adm Final (kN)', fontsize=14)
    plt.grid(True)
    plt.legend()
    plt.tight_layout()
    plt.savefig('grafico2.png', bbox_inches='tight', dpi=100)  # Reduzindo o DPI para 100

    plt.figure(figsize=(6, 8))  # Reduzindo o tamanho da figura
    plt.plot(aoki, cota, color='green', linestyle='-', linewidth=2, label='Aoki-Velloso')
    plt.scatter(aoki, cota, color='red', marker='.', linewidth=3, s=100)
    plt.axvline(cargaAdmEsperada, color='purple', linestyle='--', linewidth=3, label='Carga Adm Esperada')
    plt.axhline(0, color='k', linestyle='dotted', linewidth=2, label='N.T')
    plt.axhline(topoBloco, color='k', linestyle='dotted', linewidth=1, label='Topo do Bloco')
    plt.axhline(fundoBloco, color='k', linestyle='dotted', linewidth=1, label='Fundo do Bloco')
    plt.xlabel('Carga Admissível (kN)', fontsize=12)
    plt.ylabel('Cotas (m)', fontsize=12)
    plt.title('Carga Adm Final (kN)', fontsize=14)
    plt.grid(True)
    plt.legend()
    plt.tight_layout()
    plt.savefig('grafico3.png', bbox_inches='tight', dpi=100)  # Reduzindo o DPI para 100

    plt.figure(figsize=(6, 8))  # Reduzindo o tamanho da figura
    plt.plot(aoki, cota, color='green', linestyle='-', linewidth=2, label='Aoki-Velloso')
    plt.scatter(aoki, cota, color='red', marker='.', linewidth=3, s=100)
    plt.plot(decQuar, cota, color='blue', linestyle='-', linewidth=2, label='Décourt-Quaresma')
    plt.scatter(decQuar, cota, color='purple', marker='.', linewidth=3, s=100)
    plt.axvline(cargaAdmEsperada, color='purple', linestyle='--', linewidth=3, label='Carga Adm Esperada')
    plt.axhline(0, color='k', linestyle='dotted', linewidth=2, label='N.T')
    plt.axhline(topoBloco, color='k', linestyle='dotted', linewidth=1, label='Topo do Bloco')
    plt.axhline(fundoBloco, color='k', linestyle='dotted', linewidth=1, label='Fundo do Bloco')
    plt.xlabel('Carga Admissível (kN)', fontsize=12)
    plt.ylabel('Cotas (m)', fontsize=12)
    plt.title('Carga Admissível - Aoki e Decourt (kN)', fontsize=14)
    plt.grid(True)
    plt.legend()
    plt.tight_layout()
    plt.savefig('grafico4.png', bbox_inches='tight', dpi=100)  # Reduzindo o DPI para 100

    return 'tabela_aoki.png', 'grafico3.png', 'tabela_decQuar.png', 'grafico2.png', 'tabela_notaveis.png', 'grafico1.png', 'grafico4.png'



def resize_and_save_image(image_path, output_path, target_size=(600, 800), min_size=(50, 50), dpi=(300, 300)):
    with Image.open(image_path) as img:
        img = img.resize(target_size, resample=Image.LANCZOS)
        img.save(output_path, dpi=dpi)



class ImageWindow(QMainWindow):
    def __init__(self, image_files, image_names):
        super().__init__()

        self.image_files = image_files
        self.image_names = image_names
        self.initUI()

    def initUI(self):
        widget = QWidget()
        self.setCentralWidget(widget)

        self.scroll_area = QScrollArea(self)
        self.scroll_area.setWidgetResizable(True)

        self.image_widget = QWidget()
        self.scroll_area.setWidget(self.image_widget)

        self.layout = QVBoxLayout(self.image_widget)

        self.list_widget = QListWidget(self)
        self.list_widget.itemClicked.connect(self.onImageClicked)

        for i, (image_file, image_name) in enumerate(zip(self.image_files, self.image_names)):
            item = QListWidgetItem()
            pixmap = QPixmap(image_file)
            pixmap = pixmap.scaledToHeight(200)  # Redimensionar miniatura para altura de 200 pixels
            icon = QIcon(pixmap)  # Criar QIcon a partir do QPixmap
            item.setIcon(icon)  # Definir QIcon como ícone do item
            item.setText(image_name)  # Adicionar nome personalizado como texto do item
            self.list_widget.addItem(item)

        self.back_button = QPushButton('Back to Thumbnails', self)
        self.back_button.clicked.connect(self.onBackClicked)

        self.zoom_in_button = QPushButton('Zoom In', self)
        self.zoom_in_button.clicked.connect(self.zoomIn)

        self.zoom_out_button = QPushButton('Zoom Out', self)
        self.zoom_out_button.clicked.connect(self.zoomOut)

        layout = QVBoxLayout(widget)
        layout.addWidget(self.list_widget)
        layout.addWidget(self.back_button)
        layout.addWidget(self.zoom_in_button)
        layout.addWidget(self.zoom_out_button)

        self.scroll_area.setWidget(self.image_widget)
        layout.addWidget(self.scroll_area)

        self.setWindowTitle('Image Viewer')
        self.showMaximized()

    def onImageClicked(self, item):
        # Limpar layout e adicionar imagem em tela cheia
        self.clearLayout()
        self.current_image_file = self.image_files[self.list_widget.row(item)]
        self.pixmap = QPixmap(self.current_image_file)
        self.image_label = QLabel(self)
        self.image_label.setPixmap(self.pixmap)
        self.layout.addWidget(self.image_label)

    def clearLayout(self):
        # Remover todos os widgets do layout
        while self.layout.count():
            item = self.layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

    def onBackClicked(self):
        # Voltar ao menu de miniaturas
        self.clearLayout()
        self.initUI()

    def zoomIn(self):
        # Aumentar o fator de escala (zoom)
        if hasattr(self, 'pixmap'):
            self.pixmap = self.pixmap.scaled(self.pixmap.size() * 1.2, aspectRatioMode=Qt.KeepAspectRatio)
            self.image_label.setPixmap(self.pixmap)

    def zoomOut(self):
    # Diminuir o fator de escala (zoom)
        if hasattr(self, 'pixmap'):
            # Fator de redução
            zoom_factor = 0.8
        
            # Verificar o tamanho resultante após o zoom out
            new_size = self.pixmap.size() * zoom_factor
            min_size = QSize(50, 50)  # Tamanho mínimo desejado
        
                # Limitar o zoom out para não diminuir abaixo do tamanho mínimo
            if new_size.width() < min_size.width() or new_size.height() < min_size.height():
                return
        
            # Aplicar o zoom out à pixmap
            self.pixmap = self.pixmap.scaled(new_size, aspectRatioMode=Qt.KeepAspectRatio)
            self.image_label.setPixmap(self.pixmap)



def excel_Doc(listaTipoSolo, tipoEstaca, listaNspt, diametroEst, alturaBloco, cargaAdmEsperada, fileName):
    df_aoki = pd.DataFrame(resultAoki(listaTipoSolo, tipoEstaca, listaNspt, diametroEst))
    df_decQuar = pd.DataFrame(resulDecQuar(listaTipoSolo, tipoEstaca, listaNspt, diametroEst))
    df_notaveis = pd.DataFrame(resulNotaveis(listaTipoSolo, tipoEstaca, listaNspt, diametroEst))

    tabela_aoki, grafico3, tabela_decQuar, grafico2, tabela_notaveis, grafico1, grafico4 = geracaoGraficos(listaTipoSolo, tipoEstaca, listaNspt, diametroEst, alturaBloco, cargaAdmEsperada)
    
    # Reduzir o tamanho das imagens
    for image_path in [grafico3, grafico2, grafico1, grafico4]:
        with Image.open(image_path) as img:
            img = img.resize((600, 800), Image.LANCZOS)
            img.save(image_path)

    for image_path in [tabela_aoki, tabela_decQuar, tabela_notaveis]:
        with Image.open(image_path) as img:
            img = img.resize((800, 200), Image.LANCZOS)
            img.save(image_path)

    # Cria um novo arquivo Excel com openpyxl
    with pd.ExcelWriter(fileName + '.xlsx', engine='openpyxl') as writer:
        df_aoki.to_excel(writer, sheet_name='Resultados Aoki Velloso', index=False)
        df_decQuar.to_excel(writer, sheet_name='Resultados Decourt Quaresma', index=False)
        df_notaveis.to_excel(writer, sheet_name='Resultados Notáveis', index=False)

        # Adiciona as imagens às planilhas
        wb = writer.book

        # Formata as células para todas as planilhas
        for sheet_name in ['Resultados Aoki Velloso', 'Resultados Decourt Quaresma', 'Resultados Notáveis']:
            ws = wb[sheet_name]

            # Aplica formatação a todas as células
            for row in ws.iter_rows():
                for cell in row:
                    cell.font = Font(name='Courier New')
                    cell.alignment = Alignment(horizontal='center', vertical='center')
                    cell.border = Border(left=Side(style='thin'), 
                                         right=Side(style='thin'), 
                                         top=Side(style='thin'), 
                                         bottom=Side(style='thin'))
                 
            for i in range(1, len(ws['A'])+1):
                ws.row_dimensions[i].height = 20

            # Ajusta a altura da primeira linha (cabeçalho)
            ws.row_dimensions[1].height = 40  # Ajusta a altura para 40

            # Adiciona uma quebra de linha no cabeçalho
            for cell in ws[1]:
                cell.alignment = Alignment(wrapText=True, horizontal='center')
                cell.value = '\n' + str(cell.value)

            # Ajusta a largura das colunas
            for col in ws.columns:
                max_length = 0
                column = col[0].column_letter
                for cell in col:
                    try: 
                        if len(str(cell.value)) > max_length:
                            max_length = len(cell.value)
                    except:
                        pass
                adjusted_width = (max_length + 2)  # Adiciona 2 para um pouco de espaço extra
                ws.column_dimensions[column].width = adjusted_width

            # Adiciona as imagens
            if sheet_name == 'Resultados Aoki Velloso':
                img = xlImage(grafico3)
                ws.add_image(img, 'M1')
            elif sheet_name == 'Resultados Decourt Quaresma':
                img = xlImage(grafico2)
                ws.add_image(img, 'N1')
            elif sheet_name == 'Resultados Notáveis':
                img1 = xlImage(grafico4)
                ws.add_image(img1, 'F1')
                img2 = xlImage(grafico1)
                ws.add_image(img2, 'Q1')




def word_Doc(listaTipoSolo, tipoEstaca, listaNspt, diametroEst, alturaBloco, cargaAdmEsperada, fileName):
    df_aoki = pd.DataFrame(resultAoki(listaTipoSolo, tipoEstaca, listaNspt, diametroEst))
    df_decQuar = pd.DataFrame(resulDecQuar(listaTipoSolo, tipoEstaca, listaNspt, diametroEst))
    df_notaveis = pd.DataFrame(resulNotaveis(listaTipoSolo, tipoEstaca, listaNspt, diametroEst))

    menorAokiDec = minAokiDec(listaTipoSolo, tipoEstaca, listaNspt, diametroEst)
    cota = cotasPonta(listaTipoSolo)
    aoki = paFinalAoki(listaTipoSolo, tipoEstaca, listaNspt, diametroEst)
    decQuar = paFinalDecQuar(listaTipoSolo, tipoEstaca, listaNspt, diametroEst)

    topoBloco = -1 + alturaBloco
    fundoBloco = -1

    tabela_aoki, grafico3, tabela_decQuar, grafico2, tabela_notaveis, grafico1, grafico4 = geracaoGraficos(listaTipoSolo, tipoEstaca, listaNspt, diametroEst, alturaBloco, cargaAdmEsperada)

    template = Document()
    template.add_paragraph("Cálculo Da Capacidade de Carga De Estaca", "EstiloCabeçalho")

    for secao in template.sections:
        secao.top_margin = Cm(2)
        secao.bottom_margin = Cm(2)
        secao.left_margin = Cm(2)
        secao.right_margin = Cm(2)

    template.add_paragraph("Dados de Entrada", "EstiloTitulo")

    texto = f"""Tipo da Estaca: {tipoEstaca}
D = {diametroEst * 100} cm (Diâmetro da estaca)
Carga admissível esperada: {cargaAdmEsperada / 10} tf
Lista NSPT: {listaNspt}
Altura do bloco: {alturaBloco * 100} cm"""


    template.add_paragraph(texto, "EstiloCorpo")

    template.add_paragraph("Capacidade de carga por Aoki-Velloso", "EstiloTitulo")

    aoki_tabela = template.add_paragraph("Capacidade De Carga Por Aoki-Velloso", "EstiloCorpo")
    tabela = template.add_table(rows=df_aoki.shape[0] + 1, cols=df_aoki.shape[1], style="EstiloTabela")

    for coluna_idx, coluna_nome in enumerate(df_aoki.columns):
        cell = tabela.cell(0, coluna_idx)
        cell.text = coluna_nome

    for linha_idx in range(df_aoki.shape[0]):
        for coluna_idx in range(df_aoki.shape[1]):
            cell = tabela.cell(linha_idx + 1, coluna_idx)
            cell.text = str(df_aoki.iloc[linha_idx, coluna_idx])

    template.add_page_break()

    template.add_paragraph("Capacidade de carga por Decout-Quaresma", "EstiloTitulo")

    decQuar_tabela = template.add_paragraph("Capacidade De Carga Por Decourt-Quaresma", "EstiloCorpo")
    tabela = template.add_table(rows=df_decQuar.shape[0] + 1, cols=df_decQuar.shape[1], style="EstiloTabela")

    for coluna_idx, coluna_nome in enumerate(df_decQuar.columns):
        cell = tabela.cell(0, coluna_idx)
        cell.text = coluna_nome

    for linha_idx in range(df_decQuar.shape[0]):
        for coluna_idx in range(df_decQuar.shape[1]):
            cell = tabela.cell(linha_idx + 1, coluna_idx)
            cell.text = str(df_decQuar.iloc[linha_idx, coluna_idx])

    template.add_page_break()

    template.add_paragraph("Resultados Notáveis", "EstiloTitulo")

    notaveis_tabela = template.add_paragraph("Resultados Notáveis", "EstiloCorpo")
    tabela = template.add_table(rows=df_notaveis.shape[0] + 1, cols=df_notaveis.shape[1], style="EstiloTabela")

    for coluna_idx, coluna_nome in enumerate(df_notaveis.columns):
        cell = tabela.cell(0, coluna_idx)
        cell.text = coluna_nome

    for linha_idx in range(df_notaveis.shape[0]):
        for coluna_idx in range(df_notaveis.shape[1]):
            cell = tabela.cell(linha_idx + 1, coluna_idx)
            cell.text = str(df_notaveis.iloc[linha_idx, coluna_idx])

    template.add_page_break()

    nomeFig1 = template.add_paragraph("Carga admissível final", "EstiloTitulo")
    fig1 = template.add_paragraph()
    fig1.add_run().add_picture(grafico1, width=Inches(6))
    fig1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    template.add_page_break()

    nomeFig2 = template.add_paragraph("Comparação entre os dois métodos", "EstiloTitulo")
    fig2 = template.add_paragraph()
    fig2.add_run().add_picture(grafico4, width=Inches(6))
    fig2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    template.save(fileName + '.docx')



#print(word_Doc(listaTipoSolo, tipoEstaca, listaNspt, diametroEst, alturaBloco, cargaAdmEsperada, fileName))
#print(excel_Doc(listaTipoSolo, tipoEstaca, listaNspt, diametroEst, alturaBloco, cargaAdmEsperada, fileName))